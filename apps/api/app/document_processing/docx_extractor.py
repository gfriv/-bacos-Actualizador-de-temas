import re
from pathlib import Path

from docx import Document as DocxDocument
from docx.document import Document as DocxDocumentType
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

EMPTY_DOCX_MESSAGE = (
    "El DOCX no contiene texto extraible. Si el documento esta formado por imagenes, "
    "sera necesario OCR en una fase posterior."
)

CONTROL_CHARS_RE = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f]")
WHITESPACE_RE = re.compile(r"[ \t]+")


def extract_docx_text(path: str | Path) -> str:
    """Extract readable, section-friendly text from a DOCX.

    The MVP originally read only `document.paragraphs`, which loses tables and relevant
    header/footer context. This extractor keeps the body order for paragraphs and tables,
    converts styled headings to Markdown headings and includes non-duplicated headers/footers.
    """

    document = DocxDocument(str(path))
    parts: list[str] = []

    parts.extend(_extract_headers_and_footers(document))
    for block in _iter_block_items(document):
        if isinstance(block, Paragraph):
            rendered = _render_paragraph(block)
        else:
            rendered = _render_table(block)
        if rendered:
            parts.append(rendered)

    extracted = _join_parts(parts)
    if not extracted.strip():
        raise ValueError(EMPTY_DOCX_MESSAGE)
    return extracted


def _iter_block_items(parent: DocxDocumentType | _Cell):
    if isinstance(parent, DocxDocumentType):
        parent_element = parent.element.body
    elif isinstance(parent, _Cell):
        parent_element = parent._tc
    else:
        raise TypeError("Unsupported DOCX container.")

    for child in parent_element.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def _extract_headers_and_footers(document: DocxDocumentType) -> list[str]:
    parts: list[str] = []
    seen: set[str] = set()

    for section in document.sections:
        for label, story in (("Cabecera", section.header), ("Pie de pagina", section.footer)):
            story_parts: list[str] = []
            for paragraph in story.paragraphs:
                rendered = _render_paragraph(paragraph)
                if rendered:
                    story_parts.append(rendered)
            for table in story.tables:
                rendered = _render_table(table)
                if rendered:
                    story_parts.append(rendered)

            text = _join_parts(story_parts)
            fingerprint = _fingerprint(text)
            if text and fingerprint not in seen:
                seen.add(fingerprint)
                parts.append(f"## {label}\n\n{text}")

    return parts


def _render_paragraph(paragraph: Paragraph) -> str:
    text = _clean_text(paragraph.text)
    if not text:
        return ""

    style_name = (paragraph.style.name if paragraph.style is not None else "").strip().lower()
    heading_level = _heading_level(style_name)
    if heading_level:
        return f"{'#' * heading_level} {text}"
    if _is_list_style(style_name):
        return f"- {text}"
    return text


def _heading_level(style_name: str) -> int | None:
    matches = (
        ("heading 1", 1),
        ("titulo 1", 1),
        ("title 1", 1),
        ("heading 2", 2),
        ("titulo 2", 2),
        ("title 2", 2),
        ("heading 3", 3),
        ("titulo 3", 3),
        ("title 3", 3),
        ("heading 4", 4),
        ("titulo 4", 4),
        ("title 4", 4),
        ("title", 1),
        ("titulo", 1),
        ("subtitle", 2),
        ("subtitulo", 2),
    )
    for marker, level in matches:
        if marker in style_name:
            return level
    return None


def _is_list_style(style_name: str) -> bool:
    return any(marker in style_name for marker in ("list", "lista", "bullet", "vineta", "viñeta"))


def _render_table(table: Table) -> str:
    rows: list[list[str]] = []
    for row in table.rows:
        cells = [_cell_text(cell) for cell in row.cells]
        if any(cells):
            rows.append(cells)
    if not rows:
        return ""

    width = max(len(row) for row in rows)
    normalized_rows = [row + [""] * (width - len(row)) for row in rows]
    header = normalized_rows[0]
    separator = ["---"] * width
    body = normalized_rows[1:]
    lines = [
        "| " + " | ".join(_escape_table_cell(cell) for cell in header) + " |",
        "| " + " | ".join(separator) + " |",
    ]
    for row in body:
        lines.append("| " + " | ".join(_escape_table_cell(cell) for cell in row) + " |")
    return "\n".join(lines)


def _cell_text(cell: _Cell) -> str:
    chunks: list[str] = []
    for block in _iter_block_items(cell):
        if isinstance(block, Paragraph):
            text = _clean_text(block.text)
        else:
            text = _render_table(block)
        if text:
            chunks.append(text)
    return " ".join(chunks)


def _clean_text(value: str) -> str:
    without_controls = CONTROL_CHARS_RE.sub(" ", value)
    return WHITESPACE_RE.sub(" ", without_controls).strip()


def _escape_table_cell(value: str) -> str:
    return _clean_text(value).replace("|", "\\|")


def _join_parts(parts: list[str]) -> str:
    return "\n\n".join(part.strip() for part in parts if part and part.strip()).strip()


def _fingerprint(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip().lower()
