import re
from pathlib import Path

from docx import Document

HEADING_RE = re.compile(r"^(#{1,4})\s+(.+)$")
ORDERED_LIST_RE = re.compile(r"^\d+[\).]\s+(.+)$")
TABLE_SEPARATOR_RE = re.compile(r"^\|\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?$")


def write_markdown_docx(markdown: str, path: str | Path) -> str:
    target = Path(path)
    target.parent.mkdir(parents=True, exist_ok=True)
    document = Document()

    lines = markdown.splitlines()
    index = 0
    while index < len(lines):
        line = lines[index].rstrip()
        if not line.strip():
            index += 1
            continue

        heading = HEADING_RE.match(line)
        if heading:
            level = min(len(heading.group(1)), 4)
            document.add_heading(_clean_inline(heading.group(2)), level=level)
            index += 1
            continue

        if _is_table_line(line):
            table_lines: list[str] = []
            while index < len(lines):
                if _is_table_line(lines[index]):
                    table_lines.append(lines[index])
                    index += 1
                    continue
                if not lines[index].strip() and _next_nonempty_is_table(lines, index + 1):
                    index += 1
                    continue
                break
            _add_markdown_table(document, table_lines)
            continue

        if line.strip().startswith("- "):
            document.add_paragraph(_clean_inline(line.strip()[2:]), style="List Bullet")
            index += 1
            continue

        ordered = ORDERED_LIST_RE.match(line.strip())
        if ordered:
            document.add_paragraph(_clean_inline(ordered.group(1)), style="List Number")
            index += 1
            continue

        paragraph_lines = [line.strip()]
        index += 1
        while index < len(lines):
            next_line = lines[index].rstrip()
            if (
                not next_line.strip()
                or HEADING_RE.match(next_line)
                or _is_table_line(next_line)
                or next_line.strip().startswith("- ")
                or ORDERED_LIST_RE.match(next_line.strip())
            ):
                break
            paragraph_lines.append(next_line.strip())
            index += 1
        document.add_paragraph(_clean_inline(" ".join(paragraph_lines)))

    document.save(target)
    return str(target)


def _is_table_line(line: str) -> bool:
    stripped = line.strip()
    return stripped.startswith("|") and stripped.endswith("|") and stripped.count("|") >= 2


def _next_nonempty_is_table(lines: list[str], start_index: int) -> bool:
    for index in range(start_index, len(lines)):
        if not lines[index].strip():
            continue
        return _is_table_line(lines[index])
    return False


def _add_markdown_table(document, table_lines: list[str]) -> None:
    rows = [_split_table_row(line) for line in table_lines if not TABLE_SEPARATOR_RE.match(line.strip())]
    rows = [row for row in rows if any(cell.strip() for cell in row)]
    if not rows:
        return

    width = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=width)
    table.style = "Table Grid"
    for row_index, row in enumerate(rows):
        normalized_row = row + [""] * (width - len(row))
        for col_index, cell_text in enumerate(normalized_row):
            table.cell(row_index, col_index).text = _clean_inline(cell_text)


def _split_table_row(line: str) -> list[str]:
    stripped = line.strip().strip("|")
    return [cell.strip().replace("\\|", "|") for cell in stripped.split("|")]


def _clean_inline(text: str) -> str:
    return (
        text.replace("**", "")
        .replace("__", "")
        .replace("`", "")
        .replace("<br>", "\n")
        .replace("<br/>", "\n")
        .strip()
    )
