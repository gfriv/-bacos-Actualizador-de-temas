from io import BytesIO
from pathlib import Path

import pytest
from docx import Document
from fastapi.testclient import TestClient
from pypdf import PdfWriter

from app.core.config import settings
from app.document_processing.docx_extractor import extract_docx_text
from app.document_processing.docx_writer import write_markdown_docx
from app.document_processing.ocr_extractor import OCRConfig, OCRUnavailableError
from app.document_processing.pdf_extractor import SCANNED_PDF_MESSAGE, extract_pdf_text
from app.document_processing.section_splitter import split_sections
from tests.test_projects import create_project


def make_docx_bytes() -> bytes:
    buffer = BytesIO()
    document = Document()
    document.add_heading("1. Introducción", level=1)
    document.add_paragraph("Texto original sobre competencias, criterios y evaluación.")
    document.add_heading("2. Desarrollo", level=1)
    document.add_paragraph("Contenido científico pendiente de revisión.")
    document.save(buffer)
    buffer.seek(0)
    return buffer.read()


def make_complex_docx_bytes() -> bytes:
    buffer = BytesIO()
    document = Document()
    section = document.sections[0]
    section.header.paragraphs[0].text = "Centro de Estudios Ábacos"
    section.footer.paragraphs[0].text = "Material docente revisable"
    document.add_heading("Tema 1. Desarrollo evolutivo", level=1)
    document.add_paragraph("Texto introductorio con competencias y evaluación.")
    table = document.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "Concepto"
    table.cell(0, 1).text = "Aplicación didáctica"
    table.cell(1, 0).text = "Atención"
    table.cell(1, 1).text = "Actividades graduadas"
    table.cell(2, 0).text = "Memoria"
    table.cell(2, 1).text = "Recuperación espaciada"
    document.add_heading("1.1 Marco curricular", level=2)
    document.add_paragraph("Referencia a normativa de Extremadura para Primaria.")
    document.save(buffer)
    buffer.seek(0)
    return buffer.read()


def test_upload_document_extracts_docx(client: TestClient, auth_headers: dict[str, str]) -> None:
    project_id = create_project(client, auth_headers)
    response = client.post(
        f"/api/projects/{project_id}/documents",
        headers=auth_headers,
        files={
            "file": (
                "tema.docx",
                make_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert response.status_code == 201
    assert "Texto original" in response.json()["extracted_text"]
    assert "file_path" not in response.json()

    sections = client.get(f"/api/projects/{project_id}/sections", headers=auth_headers)
    assert sections.status_code == 200
    assert len(sections.json()) >= 1


def test_docx_extractor_preserves_headers_tables_and_headings(tmp_path: Path) -> None:
    path = tmp_path / "tema-complejo.docx"
    path.write_bytes(make_complex_docx_bytes())

    extracted = extract_docx_text(path)

    assert "Centro de Estudios Ábacos" in extracted
    assert "# Tema 1. Desarrollo evolutivo" in extracted
    assert "| Concepto | Aplicación didáctica |" in extracted
    assert "Recuperación espaciada" in extracted
    assert "## 1.1 Marco curricular" in extracted


def test_section_splitter_handles_nested_headings_and_tables() -> None:
    sections = split_sections(
        "\n".join(
            [
                "# Tema 1. Desarrollo evolutivo",
                "Texto inicial.",
                "| Concepto | Aplicación |",
                "| --- | --- |",
                "| Atención | Actividades graduadas |",
                "1.1 Marco curricular",
                "Normativa, competencias y saberes.",
            ]
        )
    )

    assert [section.title for section in sections] == [
        "Tema 1. Desarrollo evolutivo",
        "1.1 Marco curricular",
    ]
    assert "| Concepto | Aplicación |" in sections[0].content
    assert "competencias" in sections[1].detected_concepts


def test_pdf_extractor_reports_scanned_pdf(tmp_path: Path) -> None:
    path = tmp_path / "escaneado.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=72, height=72)
    with path.open("wb") as file:
        writer.write(file)

    with pytest.raises(ValueError, match=SCANNED_PDF_MESSAGE):
        extract_pdf_text(path, ocr_enabled=False)


def test_pdf_extractor_uses_ocr_for_scanned_pdf(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    path = tmp_path / "escaneado.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=72, height=72)
    with path.open("wb") as file:
        writer.write(file)

    def fake_ocr(pdf_path: str | Path, config: OCRConfig) -> str:
        assert Path(pdf_path) == path
        assert config.languages
        return "[Página 1 - OCR]\nTexto recuperado mediante OCR para análisis docente."

    monkeypatch.setattr("app.document_processing.pdf_extractor.extract_pdf_ocr_text", fake_ocr)

    extracted = extract_pdf_text(path, ocr_enabled=True)

    assert "Texto recuperado mediante OCR" in extracted


def test_pdf_extractor_reports_missing_ocr_engine(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    path = tmp_path / "escaneado.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=72, height=72)
    with path.open("wb") as file:
        writer.write(file)

    def fake_ocr(pdf_path: str | Path, config: OCRConfig) -> str:
        raise OCRUnavailableError("Tesseract OCR no está disponible en el servidor.")

    monkeypatch.setattr("app.document_processing.pdf_extractor.extract_pdf_ocr_text", fake_ocr)

    with pytest.raises(ValueError, match="Tesseract OCR no está disponible"):
        extract_pdf_text(path, ocr_enabled=True)


def test_docx_writer_exports_markdown_tables(tmp_path: Path) -> None:
    path = tmp_path / "consolidado.docx"

    write_markdown_docx(
        "\n\n".join(
            [
                "# Documento consolidado",
                "| Concepto | Aplicación |",
                "| --- | --- |",
                "| Atención | Actividades graduadas |",
                "- Revisión docente obligatoria",
            ]
        ),
        path,
    )

    document = Document(path)
    assert document.paragraphs[0].text == "Documento consolidado"
    assert len(document.tables) == 1
    assert document.tables[0].cell(1, 1).text == "Actividades graduadas"


def test_upload_invalid_docx_returns_422_without_internal_path(
    client: TestClient, auth_headers: dict[str, str]
) -> None:
    project_id = create_project(client, auth_headers)
    response = client.post(
        f"/api/projects/{project_id}/documents",
        headers=auth_headers,
        files={
            "file": (
                "tema.docx",
                b"esto no es un docx valido",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert response.status_code == 422
    assert "DOCX" in response.json()["detail"]
    assert "storage" not in response.json()["detail"].lower()


def test_database_storage_backend_serves_uploaded_document(
    client: TestClient, auth_headers: dict[str, str]
) -> None:
    original_storage_backend = settings.storage_backend
    settings.storage_backend = "database"
    try:
        project_id = create_project(client, auth_headers)
        document = client.post(
            f"/api/projects/{project_id}/documents",
            headers=auth_headers,
            files={
                "file": (
                    "tema.docx",
                    make_docx_bytes(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )
        assert document.status_code == 201

        download = client.get(f"/api/documents/{document.json()['id']}/download", headers=auth_headers)
        assert download.status_code == 200
        assert download.content.startswith(b"PK")
    finally:
        settings.storage_backend = original_storage_backend


def test_research_analysis_creates_reports_and_suggestions(
    client: TestClient, auth_headers: dict[str, str]
) -> None:
    project_id = create_project(client, auth_headers)
    client.post(
        f"/api/projects/{project_id}/documents",
        headers=auth_headers,
        files={
            "file": (
                "tema.docx",
                make_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    response = client.post(f"/api/projects/{project_id}/analysis/research", headers=auth_headers)
    assert response.status_code == 200
    assert len(response.json()["reports"]) == 6
    assert len(response.json()["suggestions"]) == 2

    reports = client.get(f"/api/projects/{project_id}/reports", headers=auth_headers)
    suggestions = client.get(f"/api/projects/{project_id}/suggestions", headers=auth_headers)
    assert len(reports.json()) == 6
    assert len(suggestions.json()) == 2
    source_report = next(report for report in reports.json() if report["report_type"] == "source_validation")
    report_download = client.get(f"/api/reports/{source_report['id']}/download", headers=auth_headers)
    assert report_download.status_code == 200
    assert b"Centro de Formaci" in report_download.content


def test_consolidation_rejects_pending_only(client: TestClient, auth_headers: dict[str, str]) -> None:
    project_id = create_project(client, auth_headers)
    client.post(
        f"/api/projects/{project_id}/documents",
        headers=auth_headers,
        files={
            "file": (
                "tema.docx",
                make_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    client.post(
        f"/api/projects/{project_id}/suggestions",
        headers=auth_headers,
        json={
            "suggestion_type": "scientific_update",
            "original_fragment": "Texto original",
            "proposed_change": "Texto actualizado",
            "justification": "Justificación docente",
            "source_reference": "MockProvider",
            "confidence_level": "medium",
        },
    )
    response = client.post(f"/api/projects/{project_id}/consolidate", headers=auth_headers)
    assert response.status_code == 400
    assert "No hay sugerencias aprobadas" in response.json()["detail"]


def test_consolidation_uses_approved_suggestions_and_generates_resource(
    client: TestClient, auth_headers: dict[str, str]
) -> None:
    project_id = create_project(client, auth_headers)
    client.post(
        f"/api/projects/{project_id}/documents",
        headers=auth_headers,
        files={
            "file": (
                "tema.docx",
                make_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    suggestion = client.post(
        f"/api/projects/{project_id}/suggestions",
        headers=auth_headers,
        json={
            "suggestion_type": "scientific_update",
            "original_fragment": "Texto original",
            "proposed_change": "Texto actualizado",
            "justification": "Justificación docente",
            "source_reference": "MockProvider",
            "confidence_level": "medium",
        },
    ).json()
    client.patch(
        f"/api/suggestions/{suggestion['id']}",
        headers=auth_headers,
        json={"status": "approved"},
    )

    consolidated = client.post(f"/api/projects/{project_id}/consolidate", headers=auth_headers)
    assert consolidated.status_code == 200
    assert "docx_path" not in consolidated.json()
    consolidated_download = client.get(
        f"/api/projects/{project_id}/consolidated/download",
        headers=auth_headers,
    )
    assert consolidated_download.status_code == 200
    assert consolidated_download.content.startswith(b"PK")
    assert "Documento consolidado" in consolidated.json()["content_markdown"]
    assert "Centro de Formación y Estudios Ábacos" in consolidated.json()["content_markdown"]
    assert "asistencia de IA" in consolidated.json()["content_markdown"]

    resource = client.post(
        f"/api/projects/{project_id}/resources",
        headers=auth_headers,
        json={"resource_type": "esquema_desarrollado"},
    )
    assert resource.status_code == 200
    assert "file_path" not in resource.json()
    resource_download = client.get(f"/api/resources/{resource.json()['id']}/download", headers=auth_headers)
    assert resource_download.status_code == 200
    assert b"Recurso" in resource_download.content
    assert "Recurso didáctico simulado" in resource.json()["content_markdown"]
    assert "Centro de Formación y Estudios Ábacos" in resource.json()["content_markdown"]
    assert "asistencia de IA" in resource.json()["content_markdown"]


def test_resource_generation_requires_consolidated_document(
    client: TestClient, auth_headers: dict[str, str]
) -> None:
    project_id = create_project(client, auth_headers)
    response = client.post(
        f"/api/projects/{project_id}/resources",
        headers=auth_headers,
        json={"resource_type": "esquema_desarrollado"},
    )
    assert response.status_code == 400
    assert "documento consolidado" in response.json()["detail"]
